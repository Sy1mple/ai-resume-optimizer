from __future__ import annotations

import base64
import re
from io import BytesIO
from textwrap import wrap

from docx import Document
from docx.shared import Inches
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models import ExportFormat, ExportRequest


MIME_TYPES = {
    ExportFormat.pdf: "application/pdf",
    ExportFormat.docx: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ExportFormat.md: "text/markdown; charset=utf-8",
    ExportFormat.txt: "text/plain; charset=utf-8",
}


def export_document(request: ExportRequest) -> tuple[bytes, str, str]:
    safe_name = _safe_filename(request.file_name or request.candidate_name or "resume")
    if request.format == ExportFormat.pdf:
        return _build_pdf(request), f"{safe_name}.pdf", MIME_TYPES[request.format]
    if request.format == ExportFormat.docx:
        return _build_docx(request), f"{safe_name}.docx", MIME_TYPES[request.format]
    if request.format == ExportFormat.md:
        return request.content.encode("utf-8"), f"{safe_name}.md", MIME_TYPES[request.format]
    return _markdown_to_plain_text(request.content).encode("utf-8"), f"{safe_name}.txt", MIME_TYPES[request.format]


def _build_pdf(request: ExportRequest) -> bytes:
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    styles = getSampleStyleSheet()
    palette = _style_palette(request.style)
    styles["Title"].textColor = colors.HexColor(palette["title"])
    styles["Heading2"].textColor = colors.HexColor(palette["accent"])
    story = []

    photo = _decode_photo(request.photo_data_url)
    if photo:
        image_buffer = BytesIO(photo)
        story.append(PdfImage(image_buffer, width=1.05 * inch, height=1.05 * inch))
        story.append(Spacer(1, 10))

    for line in request.content.splitlines():
        clean = line.strip()
        if not clean:
            story.append(Spacer(1, 7))
            continue
        if clean.startswith("# "):
            story.append(Paragraph(_escape(clean[2:]), styles["Title"]))
        elif clean.startswith("## "):
            story.append(Paragraph(_escape(clean[3:]), styles["Heading2"]))
        elif clean.startswith("- "):
            story.append(Paragraph(f"&bull; {_escape(clean[2:])}", styles["BodyText"]))
        else:
            for chunk in wrap(clean, 110) or [clean]:
                story.append(Paragraph(_escape(chunk), styles["BodyText"]))

    document.build(story)
    return buffer.getvalue()


def _build_docx(request: ExportRequest) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    photo = _decode_photo(request.photo_data_url)
    if photo:
        document.add_picture(BytesIO(photo), width=Inches(1.15))

    for line in request.content.splitlines():
        clean = line.strip()
        if not clean:
            document.add_paragraph()
        elif clean.startswith("# "):
            document.add_heading(clean[2:], level=1)
        elif clean.startswith("## "):
            document.add_heading(clean[3:], level=2)
        elif clean.startswith("- "):
            document.add_paragraph(clean[2:], style="List Bullet")
        else:
            document.add_paragraph(_markdown_to_plain_text(clean))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _decode_photo(data_url: str | None) -> bytes | None:
    if not data_url or "," not in data_url:
        return None
    try:
        raw = base64.b64decode(data_url.split(",", 1)[1])
        image = Image.open(BytesIO(raw))
        image.thumbnail((512, 512))
        output = BytesIO()
        image.convert("RGB").save(output, format="JPEG", quality=88)
        return output.getvalue()
    except Exception:
        return None


def _markdown_to_plain_text(value: str) -> str:
    text = re.sub(r"^#{1,6}\s*", "", value, flags=re.MULTILINE)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    text = re.sub(r"^\s*-\s+", "• ", text, flags=re.MULTILINE)
    return text


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", value.strip()).strip("-")
    return cleaned or "resume"


def _style_palette(style: str | None) -> dict[str, str]:
    normalized = (style or "modern").lower()
    if "executive" in normalized:
        return {"title": "#111827", "accent": "#8a5a12"}
    if "compact" in normalized:
        return {"title": "#1f2937", "accent": "#374151"}
    return {"title": "#17324d", "accent": "#216869"}


def _escape(value: str) -> str:
    return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
